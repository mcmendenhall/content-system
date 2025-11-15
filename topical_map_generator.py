# --- Import Dependencies ---
import os
import json
import openai
import numpy as np
from collections import Counter
from keybert import KeyBERT
from sentence_transformers import SentenceTransformer, util
from dotenv import load_dotenv
import asyncio
import argparse
from docx import Document

# --- Load Environment Variables ---
load_dotenv()

# --- Configuration ---
SCRAPED_DIR = "scraped-articles"
OUTPUT_DIR = "topical-maps"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
NUM_TOP_ENTITIES = 50
NUM_TOP_KEYWORDS = 30
ENTITY_GROUPING_THRESHOLD = 0.7

# (🔥 NEW) Optional Manual Hierarchy Override
MANUAL_HIERARCHY = {
    # "Parent Topic": ["Child Topic 1", "Child Topic 2"],
    # Example:
    # "Cybersecurity": ["Network Security", "Endpoint Security", "Threat Intelligence"],
}

# --- Initialize OpenAI Client (🔥 Updated for 1.0.0+)
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- Helper Functions ---

def load_scraped_articles(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    articles = data.get("articles", [])
    print(f"✅ Loaded {len(articles)} articles from {file_path}")
    return articles

async def extract_entities_openai(text):
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Extract a flat list of important named entities (people, places, organizations, products, concepts) from the text. Return ONLY a list."},
                {"role": "user", "content": text[:4000]}
            ],
            temperature=0,
            max_tokens=500
        )
        extracted = response.choices[0].message.content
        entities = [e.strip("- ").strip() for e in extracted.splitlines() if e.strip()]
        return entities
    except Exception as e:
        print(f"❌ OpenAI entity extraction error: {e}")
        return []

def extract_keywords_keybert(texts):
    model = SentenceTransformer(EMBEDDING_MODEL)
    kw_model = KeyBERT(model)
    all_keywords = []
    for text in texts:
        keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1,3), stop_words='english', top_n=10)
        all_keywords.extend([kw for kw, _ in keywords])
    print(f"✅ Extracted {len(all_keywords)} total keywords")
    return list(set(all_keywords))

def group_entities(entities):
    model = SentenceTransformer(EMBEDDING_MODEL)
    embeddings = model.encode(entities, convert_to_tensor=True)
    similarity_matrix = util.cos_sim(embeddings, embeddings)

    groups = []
    used = set()

    for idx, entity in enumerate(entities):
        if idx in used:
            continue
        group = [entity]
        used.add(idx)
        for jdx, sim_score in enumerate(similarity_matrix[idx]):
            if jdx in used:
                continue
            if sim_score > ENTITY_GROUPING_THRESHOLD:
                group.append(entities[jdx])
                used.add(jdx)
        groups.append(group)

    print(f"✅ Grouped {len(entities)} entities into {len(groups)} clusters")
    return groups

def pick_best_label(group):
    """Pick the most central entity in the group instead of random first."""
    if len(group) == 1:
        return group[0]
    model = SentenceTransformer(EMBEDDING_MODEL)
    embeddings = model.encode(group, convert_to_tensor=True)
    similarity_matrix = util.cos_sim(embeddings, embeddings)
    avg_similarity = similarity_matrix.mean(dim=1)
    best_idx = avg_similarity.argmax().item()
    return group[best_idx]

def save_topical_map(topic, grouped_entities, keywords, output_dir=OUTPUT_DIR):
    os.makedirs(output_dir, exist_ok=True)

    topical_map = {
        "main_topic": topic,
        "subtopics": [],
        "related_keywords": keywords
    }

    for group in grouped_entities:
        main_label = pick_best_label(group)
        topical_map["subtopics"].append({
            "topic": main_label,
            "related_entities": group
        })

    if MANUAL_HIERARCHY:
        topical_map["manual_hierarchy"] = MANUAL_HIERARCHY

    safe_topic = topic.strip().lower().replace(" ", "_")
    output_path = os.path.join(output_dir, f"{safe_topic}_topical_map.json")

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(topical_map, f, ensure_ascii=False, indent=2)

    print(f"✅ Saved topical map to {output_path}")

def save_topical_map_word(topic, grouped_entities, keywords, output_dir=OUTPUT_DIR):
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    doc.add_heading(f"Topical Map for '{topic}'", 0)

    doc.add_heading("Main Topic", level=1)
    doc.add_paragraph(topic)

    doc.add_heading("Subtopics", level=1)
    for group in grouped_entities:
        main_label = pick_best_label(group)
        doc.add_heading(main_label, level=2)
        for entity in group:
            doc.add_paragraph(f"- {entity}", style="List Bullet")

    if MANUAL_HIERARCHY:
        doc.add_heading("Manual Hierarchy", level=1)
        for parent, children in MANUAL_HIERARCHY.items():
            doc.add_heading(parent, level=2)
            for child in children:
                doc.add_paragraph(f"- {child}", style="List Bullet")

    doc.add_heading("Related Keywords", level=1)
    for keyword in keywords:
        doc.add_paragraph(f"- {keyword}", style="List Bullet")

    safe_topic = topic.strip().lower().replace(" ", "_")
    output_path = os.path.join(output_dir, f"{safe_topic}_topical_map.docx")
    doc.save(output_path)

    print(f"✅ Saved topical map as Word document to {output_path}")

# --- Main Processing Function ---

async def generate_topical_map(keyword):
    safe_keyword = keyword.strip().lower().replace(" ", "_")
    file_path = os.path.join(SCRAPED_DIR, f"{safe_keyword}.json")

    if not os.path.exists(file_path):
        print(f"❌ Scraped file not found: {file_path}")
        return

    print(f"🔍 Starting Topical Map Generation for: {keyword}")

    articles = load_scraped_articles(file_path)
    texts = [section.get("content", "") for article in articles for section in article.get("sections", [])]

    if not texts:
        print(f"❌ No valid article content found in {file_path}")
        return

    print("🧠 Extracting entities using OpenAI GPT-4...")
    all_entities = []
    for text in texts:
        entities = await extract_entities_openai(text)
        all_entities.extend(entities)
    top_entities = [e for e, _ in Counter(all_entities).most_common(NUM_TOP_ENTITIES)]
    print(f"✅ Found {len(top_entities)} top entities")

    print("🧠 Extracting keywords with KeyBERT...")
    keywords = extract_keywords_keybert(texts)
    keywords = keywords[:NUM_TOP_KEYWORDS]
    print(f"✅ Selected {len(keywords)} top keywords")

    print("🧠 Grouping entities into subtopics...")
    grouped_entities = group_entities(top_entities)

    print("💾 Saving topical map...")
    save_topical_map(keyword, grouped_entities, keywords)
    save_topical_map_word(keyword, grouped_entities, keywords)

    print("\n🎉 Topical Map Generation Completed Successfully!")

# --- CLI Interface ---

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Topical Map from Scraped Articles")
    parser.add_argument("keyword", type=str, help="Main keyword for the topical map")
    args = parser.parse_args()

    asyncio.run(generate_topical_map(args.keyword))